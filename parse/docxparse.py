import os
from docx import Document
import re

class Question:
	def __init__(self):
		self.question = ""
		self.power = ""
		self.answer = []
		self.prompts = []

	def __repr__(self):
		return "Power: " + self.power + " Question: " + self.question + "\n" + "Answer: " + str(self.answer) + " prompts: " + str(self.prompts) + "\n"


document = Document('source3.docx')



# for paragraph in document.paragraphs:
# 	if "page" in paragraph.text.lower() or "scop" in paragraph.text.lower() or "round" in paragraph.text.lower():
# 	 	continue
# 	# print (paragraph.text)
# 	for run in paragraph.runs:
# 		print (run.bold)
# 		print(run.text)
# 		print("+++++")

# 	print ("--------------------------------------------")
started = False
in_answer = False
curr_question = Question()
questions = []
for paragraph in document.paragraphs:
		# print (paragraph.text.lower())
	if ("tossup" in paragraph.text.lower()):
		started = True
		continue
	if any(word in paragraph.text.lower() for word in ["page", "scop", "round"]) or not started:
	 	continue
	# print (paragraph.text)
	for run in paragraph.runs:
		print (run.text)
		print (in_answer)
		if "answer: " in run.text.lower():
			print (run.text.split("Answer: "))
			curr_question.question += run.text.split("Answer: ")[0]
			if (len(run.text.split("Answer: ")) > 1):
				print ("about to add: " +str(run.text.split("Answer: ")[1]))
				curr_question.prompts.append(run.text.split("Answer: ")[1])
			print ("prompts is now "+ str(curr_question.prompts))
			in_answer = True 
			continue
		if in_answer:
			if run.bold:
				curr_question.answer.append(run.text)
			else:
				if (len(run.text) > 2):
					for word in run.text.split(" "):
						curr_question.prompts.append(word)
		else:
			if (run.bold):
				curr_question.power += run.text
			else:
				curr_question.question += run.text
	if (len(paragraph.text) == 0 and in_answer):
		questions.append(curr_question)
		curr_question = Question()
		in_answer = False
	print ("--------------------------------------------")

questions.append(curr_question)

print (questions)

print (questions[0])

print (type(questions[0].question))


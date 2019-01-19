import os
from docx import Document
import re

import sqlite3
conn = sqlite3.connect('../server/db/questions.db')
c = conn.cursor()


class Question:
	def __init__(self):
		self.question = ""
		self.power = ""
		self.answers = []
		self.prompts = []

	def __repr__(self):
		return "Power: " + self.power + " Question: " + self.question + "\n" + "Answer: " + str(self.answers) + " prompts: " + str(self.prompts) + "\n"


document = Document('source3.docx')



# for paragraph in document.paragraphs:
# 	if "page" in paragraph.text.lower() or "scop" in paragraph.text.lower() or "round" in paragraph.text.lower():
# 	 	continue
# 	# print (paragraph.text)
# 	for run in paragraph.runs:
# 		print (run.bold)
# 		print(run.text)
# 		print("+++++")

# 	print ("--------------------------------------------")
started = False
in_answer = False
curr_question = Question()
questions = []
forbidden_words = ["hsapq", "page"] #I know this filters some qs but fk it
forbidden_pattern = re.compile("^\d{3}-\d{2}-\d{2}-\d{5}$")
for paragraph in document.paragraphs:
		# print (paragraph.text.lower())
	if ("tossup" in paragraph.text.lower()):
		started = True
		continue
	if forbidden_pattern.match(paragraph.text) or any(word in paragraph.text.lower() for word in forbidden_words) or not started:
	 	continue
	# print (paragraph.text)
	for run in paragraph.runs:
		print (run.text)
		if "ANSWER: " in run.text:
			curr_question.question += run.text.split("ANSWER: ")[0]
			if (len(run.text.split("ANSWER: ")) > 1):
				curr_question.prompts.append(run.text.split("ANSWER: ")[1])
			in_answer = True 
			continue
		if in_answer:
			if run.bold:
				curr_question.answers.append(run.text)
			else:
				if (len(run.text) > 2):
					for word in run.text.split(" "):
						curr_question.prompts.append(word)
		else:
			if (run.bold):
				curr_question.power += run.text
			else:
				curr_question.question += run.text
	if (len(paragraph.text) == 0 and in_answer):
		questions.append(curr_question)
		curr_question = Question()
		in_answer = False
	print ("--------------------------------------------")

questions.append(curr_question)



def save():
	for (idx, question) in enumerate(questions):
		# c.execute("INSERT INTO questions VALUES ('2013 HSAPQ Tournament 33', idx,'question3','power4','answer4','prompt4')")
		c.execute("INSERT INTO questions VALUES (?,?,?,?)", ('2013 HSAPQ Tournament 33', idx, question.question, question.power))
		for answer in question.answers:
			print (answer)
			c.execute("INSERT INTO answers VALUES (?,?,?)", ('2013 HSAPQ Tournament 33', idx, answer))
		for prompt in question.prompts:
			print (prompt)
			c.execute("INSERT INTO prompts VALUES (?,?,?)", ('2013 HSAPQ Tournament 33', idx, prompt))

	# Save (commit) the changes
	conn.commit()

	# We can also close the connection if we are done with it.
	# Just be sure any changes have been committed or they will be lost.
	conn.close()

save()

